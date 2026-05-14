"""inspect_docx.py -- Document structure inspector for docx-read skill.

Usage:
    python inspect_docx.py <file.docx> [--format json|markdown]

Outputs paragraph counts, heading counts by level, table dimensions,
image/relationship counts, section counts, and core properties.
"""
import argparse
import json
import os
import sys

_INSTALL_CMD = (
    "python -m pip install -r .github/skills/docx-read/requirements-py38.txt"
)


def _check_deps():
    try:
        import docx  # noqa: F401
    except ImportError:
        sys.stderr.write(
            "Error: python-docx is not installed.\n"
            "Install with: {}\n".format(_INSTALL_CMD)
        )
        sys.exit(1)


def _inspect(path):
    from docx import Document
    from docx.oxml.ns import qn

    doc = Document(path)

    paragraph_count = 0
    heading_counts = {}

    for para in doc.paragraphs:
        style_name = para.style.name if para.style else ""
        if style_name.startswith("Heading"):
            parts = style_name.split()
            level = parts[-1] if parts[-1].isdigit() else "?"
            heading_counts[level] = heading_counts.get(level, 0) + 1
        else:
            paragraph_count += 1

    tables = []
    for tbl in doc.tables:
        row_count = len(tbl.rows)
        col_count = max((len(r.cells) for r in tbl.rows), default=0)
        tables.append({"rows": row_count, "cols": col_count})

    # Count image relationships
    image_count = sum(
        1 for rel in doc.part.rels.values()
        if "image" in rel.reltype
    )

    section_count = len(doc.sections)

    has_header = any(
        s.header.is_linked_to_previous is False or any(
            p.text.strip() for p in s.header.paragraphs
        )
        for s in doc.sections
    )
    has_footer = any(
        s.footer.is_linked_to_previous is False or any(
            p.text.strip() for p in s.footer.paragraphs
        )
        for s in doc.sections
    )

    props = {}
    try:
        cp = doc.core_properties
        if cp.title:
            props["title"] = cp.title
        if cp.author:
            props["author"] = cp.author
        if cp.created:
            props["created"] = cp.created.isoformat()
    except Exception:
        pass

    extras = _count_extras(path)

    return {
        "file": os.path.basename(path),
        "paragraphs": paragraph_count,
        "headings": heading_counts,
        "table_count": len(tables),
        "tables": tables,
        "images": image_count,
        "sections": section_count,
        "has_header": has_header,
        "has_footer": has_footer,
        "comment_count": extras["comment_count"],
        "tracked_insertion_count": extras["tracked_insertion_count"],
        "tracked_deletion_count": extras["tracked_deletion_count"],
        "hyperlink_count": extras["hyperlink_count"],
        "footnote_count": extras["footnote_count"],
        "endnote_count": extras["endnote_count"],
        "properties": props,
    }


def _count_extras(path):
    """Count comments, tracked changes, hyperlinks, footnotes, and endnotes.

    Read-only ZIP/XML inspection. No writing or formula evaluation.
    All counts default to 0 on any parsing error.
    """
    import zipfile
    import xml.etree.ElementTree as ET

    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    result = {
        "comment_count": 0,
        "tracked_insertion_count": 0,
        "tracked_deletion_count": 0,
        "hyperlink_count": 0,
        "footnote_count": 0,
        "endnote_count": 0,
    }
    try:
        with zipfile.ZipFile(path, "r") as z:
            names = set(z.namelist())

            if "word/comments.xml" in names:
                with z.open("word/comments.xml") as f:
                    root = ET.parse(f).getroot()
                result["comment_count"] = len(
                    root.findall("{%s}comment" % W)
                )

            if "word/document.xml" in names:
                with z.open("word/document.xml") as f:
                    root = ET.parse(f).getroot()
                result["tracked_insertion_count"] = len(
                    root.findall(".//{%s}ins" % W)
                )
                result["tracked_deletion_count"] = len(
                    root.findall(".//{%s}del" % W)
                )
                result["hyperlink_count"] = len(
                    root.findall(".//{%s}hyperlink" % W)
                )

            if "word/footnotes.xml" in names:
                with z.open("word/footnotes.xml") as f:
                    root = ET.parse(f).getroot()
                result["footnote_count"] = sum(
                    1 for fn in root.findall("{%s}footnote" % W)
                    if fn.get("{%s}type" % W) not in (
                        "separator", "continuationSeparator"
                    )
                )

            if "word/endnotes.xml" in names:
                with z.open("word/endnotes.xml") as f:
                    root = ET.parse(f).getroot()
                result["endnote_count"] = sum(
                    1 for en in root.findall("{%s}endnote" % W)
                    if en.get("{%s}type" % W) not in (
                        "separator", "continuationSeparator"
                    )
                )
    except Exception:
        pass
    return result


def _to_markdown(data):
    lines = ["# Document Inspection: {}".format(data.get("file", "")), ""]
    props = data.get("properties", {})
    if props.get("title"):
        lines.append("**Title:** {}".format(props["title"]))
    if props.get("author"):
        lines.append("**Author:** {}".format(props["author"]))
    if props.get("created"):
        lines.append("**Created:** {}".format(props["created"]))
    lines.append("")
    lines.append("| Metric | Count |")
    lines.append("|--------|-------|")
    lines.append("| Paragraphs | {} |".format(data.get("paragraphs", 0)))
    for level, count in sorted(data.get("headings", {}).items()):
        lines.append("| Heading {} | {} |".format(level, count))
    lines.append("| Tables | {} |".format(data.get("table_count", 0)))
    lines.append("| Images | {} |".format(data.get("images", 0)))
    lines.append("| Sections | {} |".format(data.get("sections", 0)))
    lines.append("| Has Header | {} |".format(data.get("has_header", False)))
    lines.append("| Has Footer | {} |".format(data.get("has_footer", False)))
    lines.append("| Comments | {} |".format(data.get("comment_count", 0)))
    lines.append("| Tracked Insertions | {} |".format(data.get("tracked_insertion_count", 0)))
    lines.append("| Tracked Deletions | {} |".format(data.get("tracked_deletion_count", 0)))
    lines.append("| Hyperlinks | {} |".format(data.get("hyperlink_count", 0)))
    lines.append("| Footnotes | {} |".format(data.get("footnote_count", 0)))
    lines.append("| Endnotes | {} |".format(data.get("endnote_count", 0)))
    if data.get("tables"):
        lines.append("")
        lines.append("## Tables")
        for i, t in enumerate(data["tables"], 1):
            lines.append("- Table {}: {}r x {}c".format(i, t["rows"], t["cols"]))
    return "\n".join(lines)


def main():
    parser = argparse.ArgumentParser(
        description="Inspect structure of a DOCX file."
    )
    parser.add_argument("path", help="Path to the .docx file")
    parser.add_argument(
        "--format", choices=["json", "markdown"], default="json",
        dest="fmt", help="Output format (default: json)"
    )
    args = parser.parse_args()

    if not os.path.exists(args.path):
        sys.stderr.write("Error: file not found: {}\n".format(args.path))
        sys.exit(1)

    _check_deps()
    data = _inspect(args.path)

    if args.fmt == "markdown":
        print(_to_markdown(data))
    else:
        print(json.dumps(data, indent=2))


if __name__ == "__main__":
    main()
