"""extract_text.py -- Content extractor for docx-read skill.

Usage:
    python extract_text.py <file.docx> [--format json|markdown]
                           [--include-tables] [--heading-only]

Outputs paragraphs, headings, and tables as sequential content blocks.
"""
import argparse
import json
import os
import sys

_INSTALL_CMD = (
    "python -m pip install -r .github/skills/docx-read/requirements-py38.txt"
)


def _check_deps():
    try:
        import docx  # noqa: F401
    except ImportError:
        sys.stderr.write(
            "Error: python-docx is not installed.\n"
            "Install with: {}\n".format(_INSTALL_CMD)
        )
        sys.exit(1)


def _extract(path, include_tables, heading_only, include_changes=False):
    from docx import Document

    doc = Document(path)
    blocks = []

    # Build a set of paragraphs that belong to tables so we can skip them
    # when iterating doc.paragraphs independently (avoid double output)
    table_para_ids = set()
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    table_para_ids.add(id(para._element))

    # Walk document body elements in order to preserve sequence
    from docx.oxml.ns import qn

    body = doc.element.body
    for child in body:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag

        if tag == "p":
            # It's a paragraph element
            from docx.text.paragraph import Paragraph
            para = Paragraph(child, doc)
            if id(child) in table_para_ids:
                continue
            style_name = para.style.name if para.style else ""
            text = para.text

            if style_name.startswith("Heading"):
                parts = style_name.split()
                level_str = parts[-1] if parts[-1].isdigit() else "1"
                level = int(level_str)
                blocks.append({
                    "type": "heading",
                    "level": level,
                    "text": text,
                })
            elif not heading_only:
                if text.strip():
                    blocks.append({
                        "type": "paragraph",
                        "text": text,
                    })

        elif tag == "tbl" and include_tables and not heading_only:
            from docx.table import Table
            tbl = Table(child, doc)
            cells = []
            for row in tbl.rows:
                row_cells = [c.text for c in row.cells]
                cells.append(row_cells)
            row_count = len(tbl.rows)
            col_count = max((len(r.cells) for r in tbl.rows), default=0)
            blocks.append({
                "type": "table",
                "rows": row_count,
                "cols": col_count,
                "cells": cells,
            })

    changes = _extract_tracked_changes(path) if include_changes else []
    return {
        "file": os.path.basename(path),
        "block_count": len(blocks),
        "blocks": blocks,
        "tracked_changes": changes,
    }


def _extract_tracked_changes(path):
    """Extract tracked insertions and deletions as text snippets.

    Read-only ZIP/XML inspection. Never modifies the file.
    Returns a list of dicts: {"type": "insertion"|"deletion", "text": "..."}.
    """
    import zipfile
    import xml.etree.ElementTree as ET

    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    results = []
    try:
        with zipfile.ZipFile(path, "r") as z:
            if "word/document.xml" not in z.namelist():
                return results
            with z.open("word/document.xml") as f:
                root = ET.parse(f).getroot()
        for ins in root.findall(".//{%s}ins" % W):
            text = "".join(t.text or "" for t in ins.findall(".//{%s}t" % W))
            if text:
                results.append({"type": "insertion", "text": text})
        for del_elem in root.findall(".//{%s}del" % W):
            text = "".join(t.text or "" for t in del_elem.findall(".//{%s}delText" % W))
            if text:
                results.append({"type": "deletion", "text": text})
    except Exception:
        pass
    return results


def _to_markdown(data):
    lines = ["# Content: {}".format(data.get("file", "")), ""]
    for block in data.get("blocks", []):
        btype = block.get("type")
        if btype == "heading":
            level = block.get("level", 1)
            lines.append("{} {}".format("#" * level, block.get("text", "")))
            lines.append("")
        elif btype == "paragraph":
            lines.append(block.get("text", ""))
            lines.append("")
        elif btype == "table":
            cells = block.get("cells", [])
            if cells:
                header = cells[0]
                lines.append("| " + " | ".join(header) + " |")
                lines.append("| " + " | ".join(["---"] * len(header)) + " |")
                for row in cells[1:]:
                    lines.append("| " + " | ".join(row) + " |")
                lines.append("")
    changes = data.get("tracked_changes", [])
    if changes:
        lines.append("## Tracked Changes")
        lines.append("")
        for ch in changes:
            prefix = "+" if ch["type"] == "insertion" else "-"
            lines.append("{} {}".format(prefix, ch["text"]))
        lines.append("")
    return "\n".join(lines)


def main():
    parser = argparse.ArgumentParser(
        description="Extract text content from a DOCX file."
    )
    parser.add_argument("path", help="Path to the .docx file")
    parser.add_argument(
        "--format", choices=["json", "markdown"], default="json",
        dest="fmt", help="Output format (default: json)"
    )
    parser.add_argument(
        "--include-tables", action="store_true",
        help="Include table content in output"
    )
    parser.add_argument(
        "--heading-only", action="store_true",
        help="Output only heading blocks"
    )
    parser.add_argument(
        "--include-changes", action="store_true",
        help="Include tracked insertions and deletions as text snippets"
    )
    args = parser.parse_args()

    if not os.path.exists(args.path):
        sys.stderr.write("Error: file not found: {}\n".format(args.path))
        sys.exit(1)

    _check_deps()
    data = _extract(args.path, args.include_tables, args.heading_only, args.include_changes)

    if args.fmt == "markdown":
        print(_to_markdown(data))
    else:
        print(json.dumps(data, indent=2))


if __name__ == "__main__":
    main()
