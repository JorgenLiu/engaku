"""inspect_docx.py -- Document structure inspector for docx-read skill.

Usage:
    python inspect_docx.py <file.docx> [--format json|markdown]

Outputs paragraph counts, heading counts by level, table dimensions,
image/relationship counts, section counts, and core properties.
"""
import argparse
import json
import os
import sys

_INSTALL_CMD = (
    "python -m pip install -r .github/skills/docx-read/requirements-py38.txt"
)


def _check_deps():
    try:
        import docx  # noqa: F401
    except ImportError:
        sys.stderr.write(
            "Error: python-docx is not installed.\n"
            "Install with: {}\n".format(_INSTALL_CMD)
        )
        sys.exit(1)


def _inspect(path):
    from docx import Document
    from docx.oxml.ns import qn

    doc = Document(path)

    paragraph_count = 0
    heading_counts = {}

    for para in doc.paragraphs:
        style_name = para.style.name if para.style else ""
        if style_name.startswith("Heading"):
            parts = style_name.split()
            level = parts[-1] if parts[-1].isdigit() else "?"
            heading_counts[level] = heading_counts.get(level, 0) + 1
        else:
            paragraph_count += 1

    tables = []
    for tbl in doc.tables:
        row_count = len(tbl.rows)
        col_count = max((len(r.cells) for r in tbl.rows), default=0)
        tables.append({"rows": row_count, "cols": col_count})

    # Count image relationships
    image_count = sum(
        1 for rel in doc.part.rels.values()
        if "image" in rel.reltype
    )

    section_count = len(doc.sections)

    has_header = any(
        s.header.is_linked_to_previous is False or any(
            p.text.strip() for p in s.header.paragraphs
        )
        for s in doc.sections
    )
    has_footer = any(
        s.footer.is_linked_to_previous is False or any(
            p.text.strip() for p in s.footer.paragraphs
        )
        for s in doc.sections
    )

    props = {}
    try:
        cp = doc.core_properties
        if cp.title:
            props["title"] = cp.title
        if cp.author:
            props["author"] = cp.author
        if cp.created:
            props["created"] = cp.created.isoformat()
    except Exception:
        pass

    return {
        "file": os.path.basename(path),
        "paragraphs": paragraph_count,
        "headings": heading_counts,
        "table_count": len(tables),
        "tables": tables,
        "images": image_count,
        "sections": section_count,
        "has_header": has_header,
        "has_footer": has_footer,
        "properties": props,
    }


def _to_markdown(data):
    lines = ["# Document Inspection: {}".format(data.get("file", "")), ""]
    props = data.get("properties", {})
    if props.get("title"):
        lines.append("**Title:** {}".format(props["title"]))
    if props.get("author"):
        lines.append("**Author:** {}".format(props["author"]))
    if props.get("created"):
        lines.append("**Created:** {}".format(props["created"]))
    lines.append("")
    lines.append("| Metric | Count |")
    lines.append("|--------|-------|")
    lines.append("| Paragraphs | {} |".format(data.get("paragraphs", 0)))
    for level, count in sorted(data.get("headings", {}).items()):
        lines.append("| Heading {} | {} |".format(level, count))
    lines.append("| Tables | {} |".format(data.get("table_count", 0)))
    lines.append("| Images | {} |".format(data.get("images", 0)))
    lines.append("| Sections | {} |".format(data.get("sections", 0)))
    lines.append("| Has Header | {} |".format(data.get("has_header", False)))
    lines.append("| Has Footer | {} |".format(data.get("has_footer", False)))
    if data.get("tables"):
        lines.append("")
        lines.append("## Tables")
        for i, t in enumerate(data["tables"], 1):
            lines.append("- Table {}: {}r x {}c".format(i, t["rows"], t["cols"]))
    return "\n".join(lines)


def main():
    parser = argparse.ArgumentParser(
        description="Inspect structure of a DOCX file."
    )
    parser.add_argument("path", help="Path to the .docx file")
    parser.add_argument(
        "--format", choices=["json", "markdown"], default="json",
        dest="fmt", help="Output format (default: json)"
    )
    args = parser.parse_args()

    if not os.path.exists(args.path):
        sys.stderr.write("Error: file not found: {}\n".format(args.path))
        sys.exit(1)

    _check_deps()
    data = _inspect(args.path)

    if args.fmt == "markdown":
        print(_to_markdown(data))
    else:
        print(json.dumps(data, indent=2))


if __name__ == "__main__":
    main()
